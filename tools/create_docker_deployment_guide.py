from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


OUTPUT = "docs/Guia_Docker_y_Despliegue_Veterinaria.docx"
PDF_OUTPUT = "docs/Guia_Docker_y_Despliegue_Veterinaria.pdf"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(35, 35, 35)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("[ ] " + item)


def add_label_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        set_cell_width(cells[0], 2700)
        set_cell_width(cells[1], 6660)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], "E8EEF5")
    doc.add_paragraph()


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    p.add_run(body)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)


def build_doc():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Guia practica: Docker y despliegue web")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Proyecto TeoNova-Vet / Veterinaria POS").italic = True

    add_callout(
        doc,
        "Recomendacion corta",
        "No es obligatorio aprender Docker para publicar la app, pero para este proyecto si es muy recomendable probarlo en Docker antes de subirlo. Te permite ensayar frontend, backend y MySQL como si estuvieran en un servidor real.",
    )

    doc.add_heading("1. Que es Docker en palabras simples", level=1)
    doc.add_paragraph(
        "Docker sirve para empacar una aplicacion con todo lo que necesita para ejecutarse: Node.js, dependencias, configuracion y servicios relacionados. En vez de decir \"en mi computador funciona\", Docker ayuda a que funcione igual en otro computador o servidor."
    )
    add_bullets(
        doc,
        [
            "Imagen: una plantilla lista para crear una app o servicio.",
            "Contenedor: una instancia corriendo de esa imagen.",
            "Docker Compose: un archivo que levanta varios contenedores juntos, por ejemplo MySQL, backend y frontend.",
            "Volumen: almacenamiento persistente para que los datos no se borren al reiniciar contenedores.",
        ],
    )

    doc.add_heading("2. Como aplica a este proyecto", level=1)
    add_label_table(
        doc,
        [
            ("Frontend", "React + Vite. Se compila con npm run build y genera archivos estaticos en dist."),
            ("Backend", "NestJS. Se compila con npm run build y arranca con node dist/main."),
            ("Base de datos", "MySQL. El esquema esta en la carpeta database."),
            ("API", "El frontend llama a /api, por eso en produccion conviene usar un proxy que envie /api al backend."),
        ],
    )

    doc.add_heading("3. Que debes instalar en Windows", level=1)
    add_numbered(
        doc,
        [
            "Instala Docker Desktop para Windows desde el sitio oficial de Docker.",
            "Durante la instalacion, permite usar WSL 2 si Windows lo solicita.",
            "Reinicia el computador si el instalador lo pide.",
            "Abre Docker Desktop y espera a que diga que esta corriendo.",
            "Abre PowerShell y verifica la instalacion.",
        ],
    )
    add_code(doc, "docker --version\ndocker compose version")

    doc.add_heading("4. Comandos basicos que vas a usar", level=1)
    add_label_table(
        doc,
        [
            ("docker ps", "Ver contenedores corriendo."),
            ("docker ps -a", "Ver todos los contenedores, incluso detenidos."),
            ("docker compose up", "Levantar los servicios definidos en docker-compose.yml."),
            ("docker compose up --build", "Reconstruir imagenes y levantar servicios."),
            ("docker compose down", "Detener y eliminar contenedores de ese proyecto."),
            ("docker compose logs -f backend", "Ver logs del backend en vivo."),
        ],
    )

    doc.add_heading("5. Ruta recomendada para tu app", level=1)
    add_numbered(
        doc,
        [
            "Primero confirma que la app compila localmente: backend y frontend.",
            "Despues crea Dockerfile para backend.",
            "Crea Dockerfile para frontend con Nginx.",
            "Crea docker-compose.yml con mysql, backend y frontend.",
            "Carga la base de datos usando database/schema.sql o la migracion correcta.",
            "Prueba login, ventas, inventario, clientes, reportes y exportacion Excel.",
            "Cuando todo funcione en Docker local, subelo a un VPS con Docker.",
        ],
    )

    doc.add_heading("6. Variables importantes de produccion", level=1)
    add_label_table(
        doc,
        [
            ("NODE_ENV", "production"),
            ("PORT", "3001 o el puerto interno del backend"),
            ("DB_HOST", "mysql si usas docker-compose, o host de la base si es externa"),
            ("DB_NAME", "veterinaria_pos u otro nombre elegido"),
            ("DB_USERNAME", "usuario de MySQL"),
            ("DB_PASSWORD", "contrasena segura"),
            ("JWT_SECRET", "cadena larga, aleatoria y secreta"),
            ("CORS_ORIGINS", "dominio real, por ejemplo https://tudominio.com"),
            ("TZ", "America/Bogota"),
            ("DB_TIMEZONE", "-05:00"),
        ],
    )

    doc.add_heading("7. Checklist antes de subir a internet", level=1)
    add_checklist(
        doc,
        [
            "El backend compila sin errores.",
            "El frontend compila sin errores.",
            "El backend arranca con NODE_ENV=production.",
            "La base de datos esta creada y tiene tablas.",
            "El usuario administrador existe.",
            "JWT_SECRET no es el valor de ejemplo.",
            "CORS_ORIGINS tiene el dominio correcto.",
            "El frontend puede entrar a /api sin error de CORS.",
            "Hay respaldo de la base de datos.",
            "El dominio tiene HTTPS activo.",
        ],
    )

    doc.add_heading("8. Donde hospedarlo", level=1)
    doc.add_paragraph(
        "Para este proyecto, mi recomendacion principal es un VPS con Docker. No es la unica opcion, pero es la mas clara para controlar backend, frontend y MySQL sin depender de paneles limitados."
    )
    add_label_table(
        doc,
        [
            ("VPS con Docker", "Recomendado. Bueno para Hostinger VPS, DigitalOcean, Contabo, Hetzner o AWS Lightsail."),
            ("Railway / Render", "Mas simple al inicio, pero revisa costos y soporte de MySQL."),
            ("Hosting compartido", "Solo si soporta Node.js de verdad. Puede ser incomodo para NestJS."),
            ("Vercel/Netlify", "Sirven para frontend, pero necesitarias backend y base de datos aparte."),
        ],
    )

    doc.add_heading("9. Plan de aprendizaje rapido", level=1)
    add_numbered(
        doc,
        [
            "Dia 1: instalar Docker Desktop y correr docker --version.",
            "Dia 2: levantar un MySQL de prueba con Docker.",
            "Dia 3: crear Dockerfile del backend.",
            "Dia 4: crear Dockerfile del frontend.",
            "Dia 5: unir todo con docker-compose.",
            "Dia 6: probar flujo completo de la app.",
            "Dia 7: preparar VPS y subir primera version.",
        ],
    )

    doc.add_heading("10. Siguiente paso recomendado", level=1)
    doc.add_paragraph(
        "El siguiente paso concreto seria crear los archivos Dockerfile y docker-compose.yml dentro del proyecto. Con eso podras ejecutar un solo comando para levantar la app completa:"
    )
    add_code(doc, "docker compose up --build")
    doc.add_paragraph(
        "Despues de eso, la misma estructura se puede usar como base para publicar en un VPS."
    )

    doc.save(OUTPUT)


def p(text, style):
    return Paragraph(text.replace("\n", "<br/>"), style)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="GuideTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideH1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideH2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=10,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="GuideCode",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=8.5,
            leading=11,
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#DADCE0"),
            borderWidth=0.5,
            borderPadding=6,
            spaceAfter=8,
        )
    )

    story = []
    story.append(p("Guia practica: Docker y despliegue web", styles["GuideTitle"]))
    story.append(p("<i>Proyecto TeoNova-Vet / Veterinaria POS</i>", styles["Body"]))

    def callout(title, body):
        table = Table(
            [[p(f"<b>{title}:</b> {body}", styles["Body"])]],
            colWidths=[6.5 * inch],
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.extend([table, Spacer(1, 10)])

    def bullets(items):
        for item in items:
            story.append(p(f"&bull; {item}", styles["Body"]))

    def numbers(items):
        for idx, item in enumerate(items, 1):
            story.append(p(f"{idx}. {item}", styles["Body"]))

    def label_table(rows):
        table_rows = [
            [
                p(f"<b>{label}</b>", styles["Body"]),
                p(detail, styles["Body"]),
            ]
            for label, detail in rows
        ]
        table = Table(table_rows, colWidths=[1.85 * inch, 4.65 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D3E0")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([table, Spacer(1, 10)])

    callout(
        "Recomendacion corta",
        "No es obligatorio aprender Docker para publicar la app, pero para este proyecto si es muy recomendable probarlo en Docker antes de subirlo. Te permite ensayar frontend, backend y MySQL como si estuvieran en un servidor real.",
    )

    story.append(p("1. Que es Docker en palabras simples", styles["GuideH1"]))
    story.append(
        p(
            'Docker sirve para empacar una aplicacion con todo lo que necesita para ejecutarse. En vez de decir "en mi computador funciona", Docker ayuda a que funcione igual en otro computador o servidor.',
            styles["Body"],
        )
    )
    bullets(
        [
            "Imagen: una plantilla lista para crear una app o servicio.",
            "Contenedor: una instancia corriendo de esa imagen.",
            "Docker Compose: un archivo que levanta varios contenedores juntos.",
            "Volumen: almacenamiento persistente para que los datos no se borren al reiniciar contenedores.",
        ]
    )

    story.append(p("2. Como aplica a este proyecto", styles["GuideH1"]))
    label_table(
        [
            ("Frontend", "React + Vite. Se compila con npm run build y genera archivos estaticos en dist."),
            ("Backend", "NestJS. Se compila con npm run build y arranca con node dist/main."),
            ("Base de datos", "MySQL. El esquema esta en la carpeta database."),
            ("API", "El frontend llama a /api, por eso en produccion conviene usar un proxy hacia el backend."),
        ]
    )

    story.append(p("3. Que debes instalar en Windows", styles["GuideH1"]))
    numbers(
        [
            "Instala Docker Desktop para Windows desde el sitio oficial de Docker.",
            "Permite usar WSL 2 si Windows lo solicita.",
            "Reinicia el computador si el instalador lo pide.",
            "Abre Docker Desktop y espera a que diga que esta corriendo.",
            "Abre PowerShell y verifica la instalacion.",
        ]
    )
    story.append(p("docker --version\ndocker compose version", styles["GuideCode"]))

    story.append(p("4. Comandos basicos que vas a usar", styles["GuideH1"]))
    label_table(
        [
            ("docker ps", "Ver contenedores corriendo."),
            ("docker ps -a", "Ver todos los contenedores, incluso detenidos."),
            ("docker compose up", "Levantar los servicios definidos en docker-compose.yml."),
            ("docker compose up --build", "Reconstruir imagenes y levantar servicios."),
            ("docker compose down", "Detener y eliminar contenedores de ese proyecto."),
            ("docker compose logs -f backend", "Ver logs del backend en vivo."),
        ]
    )

    story.append(p("5. Ruta recomendada para tu app", styles["GuideH1"]))
    numbers(
        [
            "Confirma que backend y frontend compilan localmente.",
            "Crea Dockerfile para backend.",
            "Crea Dockerfile para frontend con Nginx.",
            "Crea docker-compose.yml con mysql, backend y frontend.",
            "Carga la base de datos usando database/schema.sql o la migracion correcta.",
            "Prueba login, ventas, inventario, clientes, reportes y exportacion Excel.",
            "Cuando todo funcione en Docker local, subelo a un VPS con Docker.",
        ]
    )

    story.append(p("6. Variables importantes de produccion", styles["GuideH1"]))
    label_table(
        [
            ("NODE_ENV", "production"),
            ("PORT", "3001 o el puerto interno del backend"),
            ("DB_HOST", "mysql si usas docker-compose, o host de la base si es externa"),
            ("DB_NAME", "veterinaria_pos u otro nombre elegido"),
            ("DB_USERNAME", "usuario de MySQL"),
            ("DB_PASSWORD", "contrasena segura"),
            ("JWT_SECRET", "cadena larga, aleatoria y secreta"),
            ("CORS_ORIGINS", "dominio real, por ejemplo https://tudominio.com"),
            ("TZ", "America/Bogota"),
            ("DB_TIMEZONE", "-05:00"),
        ]
    )

    story.append(p("7. Checklist antes de subir a internet", styles["GuideH1"]))
    bullets(
        [
            "[ ] El backend compila sin errores.",
            "[ ] El frontend compila sin errores.",
            "[ ] El backend arranca con NODE_ENV=production.",
            "[ ] La base de datos esta creada y tiene tablas.",
            "[ ] El usuario administrador existe.",
            "[ ] JWT_SECRET no es el valor de ejemplo.",
            "[ ] CORS_ORIGINS tiene el dominio correcto.",
            "[ ] El frontend puede entrar a /api sin error de CORS.",
            "[ ] Hay respaldo de la base de datos.",
            "[ ] El dominio tiene HTTPS activo.",
        ]
    )

    story.append(p("8. Donde hospedarlo", styles["GuideH1"]))
    story.append(
        p(
            "Para este proyecto, mi recomendacion principal es un VPS con Docker. No es la unica opcion, pero es la mas clara para controlar backend, frontend y MySQL sin depender de paneles limitados.",
            styles["Body"],
        )
    )
    label_table(
        [
            ("VPS con Docker", "Recomendado. Bueno para Hostinger VPS, DigitalOcean, Contabo, Hetzner o AWS Lightsail."),
            ("Railway / Render", "Mas simple al inicio, pero revisa costos y soporte de MySQL."),
            ("Hosting compartido", "Solo si soporta Node.js de verdad. Puede ser incomodo para NestJS."),
            ("Vercel/Netlify", "Sirven para frontend, pero necesitarias backend y base de datos aparte."),
        ]
    )

    story.append(p("9. Plan de aprendizaje rapido", styles["GuideH1"]))
    numbers(
        [
            "Dia 1: instalar Docker Desktop y correr docker --version.",
            "Dia 2: levantar un MySQL de prueba con Docker.",
            "Dia 3: crear Dockerfile del backend.",
            "Dia 4: crear Dockerfile del frontend.",
            "Dia 5: unir todo con docker-compose.",
            "Dia 6: probar flujo completo de la app.",
            "Dia 7: preparar VPS y subir primera version.",
        ]
    )

    story.append(p("10. Siguiente paso recomendado", styles["GuideH1"]))
    story.append(
        p(
            "El siguiente paso concreto seria crear los archivos Dockerfile y docker-compose.yml dentro del proyecto. Con eso podras ejecutar un solo comando para levantar la app completa:",
            styles["Body"],
        )
    )
    story.append(p("docker compose up --build", styles["GuideCode"]))

    doc = SimpleDocTemplate(
        PDF_OUTPUT,
        pagesize=letter,
        rightMargin=1 * inch,
        leftMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        title="Guia Docker y Despliegue Veterinaria",
    )
    doc.build(story)


if __name__ == "__main__":
    build_doc()
    build_pdf()
